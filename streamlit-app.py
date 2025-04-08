import streamlit as st
import pandas as pd
import docx
import io
import requests
import PyPDF2


MODEL = "llama3"

st.set_page_config(page_title="The Fosse Fantastic File Processor", layout="centered")
st.title(f"🦙✨ The Fosse Fantastic File Processor with {MODEL}")

uploaded_file = st.file_uploader("📤 Upload a file (.xlsx, .docx, .pdf)", type=["xlsx", "docx", "pdf"])
instruction = st.text_area("🧠 Describe what you'd like done to the file")

# Session state to allow reprocessing
if "last_response" not in st.session_state:
    st.session_state.last_response = None

if uploaded_file:
    file_type = uploaded_file.name.split('.')[-1]
    if file_type == "xlsx":
        df = pd.read_excel(uploaded_file)
        st.subheader("📊 Excel Preview")
        st.dataframe(df.head())
    elif file_type == "docx":
        doc = docx.Document(uploaded_file)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        st.subheader("📄 Word Preview")
        st.text_area("Text from Document", full_text[:1000], height=300)
    elif file_type == "pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        pdf_text = ""
        for page in pdf_reader.pages:
            pdf_text += page.extract_text() or ""

        st.subheader("📕 PDF Preview")
        st.text_area("Extracted PDF Text", pdf_text[:1000], height=300)

if uploaded_file and instruction:

    def llama_process(prompt_text):
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": MODEL, "prompt": prompt_text, "stream": False}
        )
        return response.json()["response"]

    if file_type == "xlsx":
        csv_data = df.to_csv(index=False)
        prompt = f"""You are a helpful data assistant.
Here is an Excel table as CSV:

{csv_data}

Instruction:
{instruction}

Return only the updated table as CSV with headers, no explanation.
"""

        with st.spinner(f"Processing Excel with {MODEL}..."):
            result = llama_process(prompt)

        try:
            result_df = pd.read_csv(io.StringIO(result))
        except Exception:
            st.error(f"❌ Could not parse {MODEL} response as CSV")
            st.code(result)
            st.stop()

        st.success("✅ Processed Excel Output")
        st.subheader("🔍 Summary")
        st.dataframe(result_df.head())
        st.caption(f"{len(result_df)} rows × {len(result_df.columns)} columns")

        output = io.BytesIO()
        result_df.to_excel(output, index=False, engine='openpyxl')
        output.seek(0)

        st.download_button("📥 Download Excel", data=output, file_name="processed.xlsx")
        st.session_state.last_response = result_df

    elif file_type == "docx":
        prompt = f"""You are a smart writing assistant.
Here is the content of a Word document:

{full_text}

Instruction:
{instruction}

Return the updated document text only, no explanation.
"""

        with st.spinner(f"Processing Word with {MODEL}..."):
            updated_text = llama_process(prompt)

        st.subheader("🔍 Summary of Updated Word Doc")
        st.text_area("Preview", updated_text[:1000], height=300)
        st.caption(f"Total characters: {len(updated_text)}")

        new_doc = docx.Document()
        for para in updated_text.strip().split("\n"):
            new_doc.add_paragraph(para)

        docx_io = io.BytesIO()
        new_doc.save(docx_io)
        docx_io.seek(0)

        st.download_button("📥 Download Word", data=docx_io, file_name="processed.docx")
        st.session_state.last_response = updated_text

    elif file_type == "pdf":
        prompt = f"""You are a PDF assistant.
Here is the content extracted from a PDF:

{pdf_text}

Instruction:
{instruction}

Return the updated or summarized text only, no explanation.
"""

        with st.spinner(f"Processing PDF with {MODEL}..."):
            pdf_result = llama_process(prompt)

        st.subheader("🔍 Processed PDF Output")
        st.text_area("Preview", pdf_result[:1000], height=300)
        st.caption(f"Total characters: {len(pdf_result)}")

        st.download_button("📥 Download Processed Text", data=pdf_result, file_name="processed.txt")
        st.session_state.last_response = pdf_result

# 🔁 Process Again Button
if st.session_state.last_response and len(st.session_state.last_response)> 0:
    st.markdown("---")
    st.subheader("🔁 Want to try again?")
    new_instruction = st.text_area("Enter a new instruction and reprocess:", key="reprocess_input")
    if st.button("Process Again"):
        instruction = new_instruction
        uploaded_file = uploaded_file  # trigger rerun
        st.rerun()
